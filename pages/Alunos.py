import streamlit as st
from io import BytesIO
# from background import getImage, remove_bg
from docx import Document
import datetime

import firebase_admin
from firebase_admin import credentials, db
from firebase_admin import storage

firebase_admin.get_app()

def create_word_document(file_path, nome, cpf, cep, num_casa, curso, periodo, identidade_genero, identidade_racial, institutos):
    doc = Document()

    doc.add_heading('Contrato de Matrícula', level=1)

    doc.add_paragraph(f"Nome Completo: {nome}")
    doc.add_paragraph(f"CPF: {cpf}")
    doc.add_paragraph(f"CEP: {cep}")
    doc.add_paragraph(f"Nº da Casa: {num_casa}")
    doc.add_paragraph(f"Curso: {curso}")
    doc.add_paragraph(f"Período: {periodo}")
    doc.add_paragraph(f"Identidade de Gênero: {identidade_genero}")
    doc.add_paragraph(f"Identidade Racial: {identidade_racial}")
    doc.add_paragraph(f"Institutos de Interesse: {', '.join(institutos)}")

    doc.add_heading('Termos do Contrato', level=2)
    doc.add_paragraph('Termo 1: ...')
    doc.add_paragraph('Termo 2: ...')
    doc.add_page_break()
    doc.add_paragraph('Assinatura do Aluno: ________________________')
    doc.add_paragraph('Data: ___/___/____')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def upload_photo_to_storage(photo_bytes, student_id):
    firebase_admin.get_app()

    bucket = storage.bucket('herbert2024-be557.appspot.com')

    blob = bucket.blob(f'student_photos/{student_id}.jpg')

    blob.upload_from_string(photo_bytes.getvalue(), content_type='image/jpeg')

    blob.make_public()

    return blob.public_url

def register_student_to_firebase(nome, cpf, nascimento, telefone, email, cep, num_casa, curso, periodo, genero, racial, instituicoes, photo=None):
    data_atual = datetime.date.today()
    ano_atual = data_atual.year
    # Verifica se o CPF já está cadastrado
    
    firebase_admin.get_app()

    
    ref = db.reference('/students')
    students = ref.order_by_child('cpf').equal_to(cpf).get()

    if students is None:
        return False  # CPF já existe
    
    student_data = {
        'nome': nome,
        'cpf': cpf,
        'nascimento': str(nascimento),
        'telefone': telefone,
        'email': email,
        'cep': cep,
        'num_casa': num_casa,
        'curso': curso,
        'periodo': periodo,
        'genero': genero,
        'etnia': racial,
        'instituicoes': instituicoes,
        'ano': ano_atual,
        'frequencia': {
            "Matemática": 100,
            "Física": 100,
            "Química": 100,
            "Biologia": 100,
            "História": 100,
            "Geografia": 100,
            "Filosofia": 100,
            "Sociologia": 100,
            "Gramática": 100,
            "Literatura": 100,
            "Redação": 100
        },
        'score': {
            "Matemática": 0,
            "Física": 0,
            "Química": 0,
            "Biologia": 0,
            "História": 0,
            "Geografia": 0,
            "Filosofia": 0,
            "Sociologia": 0,
            "Gramática": 0,
            "Literatura": 0,
            "Redação": 0
        },
        'simulados': {
            "Unicamp": [],
            "USP": [],
            "Unesp": [],
            "Enem": [],
            "Cotuca": [],
            "Etec": [],
            "IF": [],
        }
    }

    # Envia os dados para o Firebase
    new_student_ref = ref.push(student_data)

    photo_url = upload_photo_to_storage(photo, cpf)  # Usando CPF como identificador único
    student_data['photo'] = photo_url
    new_student_ref.update({'photo': photo_url})

    return True

st.title("Cadastro Herbert")

st.subheader("Dados Pessoais")

nome = st.text_input("Nome Completo: ")

col1, col2, col3 = st.columns(3)

with col1:
    cpf = st.text_input("CPF:", max_chars=11)
    
with col2:
    nascimento = st.date_input("Data de Nascimento: ", format="DD/MM/YYYY")

with col3:
    telefone = st.text_input("Whatsapp: ")

email = st.text_input("Email: ")

st.subheader("Dados Residênciais: ")

col4, col5 = st.columns(2)

with col4:
    cep = st.text_input("CEP:", max_chars=9)

with col5:
    num_casa = st.number_input("Nº: ", min_value=0, max_value=9999)    

st.subheader("Dados Herbert")

col6, col7 = st.columns(2)

with col6:
    curso = st.selectbox("Curso", ["Pré-Vestibular", "Pré-Técnico", "Concurso Público"])

periodos = ["Selecione o Período", "Matutino", "Vespertino", "Noturno"]

if curso == "Pré-Vestibular":
    periodos = ["Selecione o Período", "Matutino", "Vespertino", "Noturno"]
elif curso == "Pré-Técnico":
    periodos = ["Selecione o Período", "Matutino", "Vespertino", "Sábado"]
elif curso == "Concurso Público":
    periodos = ["Sábado"]


with col7:
    periodo = st.selectbox("Período", periodos)

instituicoes = st.multiselect("Quais institutos pretende prestar?", ["Unicamp", "USP", "Cotuca", "ETEC", "Instituto Federal", "Unesp", "Enem", "Cotil", "Particulares"])

st.subheader("Identificação")

photo = st.camera_input("Carometro")

if photo is not None:
    photo_bytes = BytesIO(photo.getvalue())
    # processed_image = remove_bg(photo_bytes)

st.subheader("Informações CAD")

col4, col5 = st.columns(2)

with col4:
    genero = st.radio("**Como você se identifica**", ["Masculino", "Feminino", "Não Binário", "Outros"])

with col5:
    racial = st.radio("Como você se identifica", ["Negro/Negra", "Branco/Branca", "Pardo/Parda", "Amarelo/Amarela", "Indígena"])

st.subheader("Informações Socieconômicas")

st.selectbox("Como se locomove para o Herbert?", ["Ônibus ou Transporte Público", "Carro", "Andando", "Outros"])

st.selectbox("Qual sua situação escolar?", ["Escola Pública", "Escola Particular", "Outro", "Outros"])


cadastrar = st.button("Cadastrar")

if cadastrar:

    registrar = register_student_to_firebase(nome, cpf, nascimento, telefone, email, cep, num_casa, curso, periodo, genero, racial, instituicoes, photo_bytes)

    if registrar == True:
            
        docx_buffer = create_word_document(".", nome, cpf, cep, num_casa, curso, periodo, genero, racial, instituicoes)

        st.download_button(label="Baixar Contrato",
                        data=docx_buffer,
                        file_name="contrato.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


        st.success("Cadastrado com sucesso!")
    else:
        st.error("Não cadastrado. Erro.")